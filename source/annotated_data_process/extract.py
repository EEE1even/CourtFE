import requests
import pandas as pd
import docx
import os
from tqdm import tqdm
from loguru import logger

extract_prompt = """

"""

def docx_extract(text):
    url = 'http://172.25.34.208:8000/v1/chat/completions'
    prompt = extract_prompt
    data = {
        "model": "qwen",
        "messages": [
            {"role": "system","content":prompt},
            {"role": "user", "content": text}
        ],
        "temperature": 0.5,
        "n": 1,
        "max_tokens": 4096,
        "stop": [
            "<|im_end|>",
            "<|endoftext|>"
        ]
    }
    response = requests.post(url,json=data)
    if response.status_code == 200:
        result = response.json()['choices'][0]['message']['content']
        return result
    else:
        return None

def docx_extract_top3(text):
    url = 'http://172.25.34.208:8000/v1/chat/completions'
    prompt = extract_prompt
    data = {
        "model": "qwen",
        "messages": [
            {"role": "system","content":prompt},
            {"role": "user", "content": text}
        ],
        "temperature": 0.8,
        "n": 3,
        "max_tokens": 4096,
        "stop": [
            "<|im_end|>",
            "<|endoftext|>"
        ]
    }
    response = requests.post(url,json=data)
    if response.status_code == 200:
        result = response.json()['choices']
        return result
    else:
        return None

def save_text_to_file(text,file_path):
    # if not text:
    #     logger.info('text is empty, jump write process.')
    #     return

    with open(file_path,'a',encoding='utf-8') as f:
        f.write(text + '\n')
        logger.info(f'{file_path} write success！')

def main(main_path,if_top3):
    docx_entries = [
        entry for entry in os.scandir(main_path)
        if entry.is_file() and entry.name.endswith('.docx')
    ]

    text_len = 0
    nums = 0

    for entry in tqdm(docx_entries,desc='Processing',total=len(docx_entries)):
        doc_path = entry.path
        names = entry.name.split('.docx')[0]
        names = names + '-ie.txt'
        # 处理未存在的路径

        text_context = ''
        doc = docx.Document(doc_path)
        # for para in doc.paragraphs:
        #     text_context.append(para.text)
        text_context = '\n'.join(para.text for para in doc.paragraphs)
        text_len += len(text_context)

        if if_top3:
            logger.info(f"{names} Top3 Processing")
            s_path = '/home/hsy/AI_arbitration/extract/data_top3/' + names
            if not os.path.exists(s_path):
                result = docx_extract_top3(text_context)
                if result != None:
                    for i in range(len(result)):
                        text = result[i]['message']['content']
                        save_text_to_file(text,s_path)
            else:
                logger.info(f'{s_path} have processed ! continue...')
        else:
            logger.info(f"{names} Processing")
            s_path = '/home/hsy/AI_arbitration/extract/data/' + names
            if not os.path.exists(s_path):
                result = docx_extract(text_context)
                if result != None:
                    logger.info(f"当前文本长度:{len(text_context)}")
                    save_text_to_file(result,s_path)
                    nums += 1
                else:
                    logger.info(f'{names}抽取长度超出处理长度')
            else:
                logger.info(f'{s_path} exist ! continue...')
                continue
        logger.info(f'{doc_path} process success !')
    mean_len = text_len / nums
    logger.info(f"抽取文本的平均长度为:{mean_len}")

if __name__ == '__main__':
    logger.add("./ie_V01.log")
    main('/home/hsy/AI_arbitration/data/docx',False)