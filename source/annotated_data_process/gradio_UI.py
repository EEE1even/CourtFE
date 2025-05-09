import json
import docx
import re
import os
import gradio as gr
import time

def read_data(path):
    """高效读取事实认定数据"""
    pattern = r'"事实认定"\s*:\s*"([^"]+)"'
    with open(path, 'r', encoding='utf-8') as f:
        return re.findall(pattern, f.read())

def load_unprocessed_files(output_path, existing_data):
    """加载待处理案件文件"""
    tingshen_path = '/home/tmpuser/AI_arbitration/data/docx'
    extract_facts_path = '/home/tmpuser/dingke/data_top3'
    
    # 类型安全处理existing_data
    processed_ids = set()
    if isinstance(existing_data, list):
        processed_ids = {item["卷宗编号"] for item in existing_data if isinstance(item, dict)}
    
    unprocessed = []
    for entry in os.scandir(tingshen_path):
        if not entry.name.endswith('.docx'):
            continue
        
        base_name = entry.name.replace('.docx', '')
        ts_id = base_name.split('-ts')[0]
        
        if ts_id in processed_ids:
            continue

        fact_path = os.path.join(extract_facts_path, f"{base_name}-ie.txt")
        if not os.path.exists(fact_path):
            continue

        try:
            doc = docx.Document(entry.path)
            transcript = doc.paragraphs[0].text if doc.paragraphs else ""
            facts = read_data(fact_path)
            fact_count = len(facts)
            
            if facts:
                unprocessed.append({
                    "id": ts_id,
                    "transcript": transcript,
                    "facts": facts,
                    "fact_count": fact_count,  # 新增字段，表示事实数量
                    "docx_path": entry.path,
                    "fact_path": fact_path
                })
        except Exception as e:
            print(f"文件处理错误: {entry.path} - {str(e)}")

    return unprocessed

def create_interface(output_path):
    """创建增强型标注界面"""
    with gr.Blocks(title="案件事实标注系统", css="""
        .fact-container {max-height: 400px; overflow-y: auto;}
        .panel {padding: 10px !important;}
    """) as interface:
        # 状态管理
        current_case = gr.State()
        case_queue = gr.State([])
        added_facts = gr.State([])
        dataset = gr.State([])

        # 界面组件
        with gr.Row():
            with gr.Column(scale=3):
                gr.Markdown("### 案件信息")
                case_id = gr.Textbox(label="卷宗编号", interactive=False)
                transcript = gr.Textbox(label="庭审笔录", lines=25, 
                                      interactive=False, show_copy_button=True)
            
            # =============================修改模块=========================
            with gr.Column(scale=2):
                gr.Markdown("### 标注操作")
                progress = gr.Markdown("系统初始化中...")
                

                fact_panel = gr.Column(visible=False)
                with fact_panel:
                    
                    # fact_checkboxes = [gr.Checkbox(label=f"事实 {i+1}", value=False) for i in range(10)]
                    # fact_textboxes = [gr.Textbox(show_label=False, lines=1) for _ in range(10)]
                    # for i in range(10):  # 假设最多 10 个事实条目
                    #     with gr.Row():
                    #         fact_checkboxes.append(gr.Checkbox(label=f"事实 {i+1}", value=False))
                    #         fact_textboxes.append(gr.Textbox(interactive=True, lines=1, show_label=False))
                      
                    fact_checkboxes=[]
                    fact_textboxes=[]
                with gr.Row():
                    submit_btn = gr.Button("提交标注", variant="primary")
                    refresh_btn = gr.Button("刷新列表")
                status = gr.Markdown("系统状态：就绪")

                # 需求：模块变大
                with gr.Column():
                    new_fact_input = gr.Textbox(label="新增事实内容", placeholder="输入新的事实认定...",lines=5,scale=10)
                    add_btn = gr.Button("添加事实", variant="secondary")
                hint = gr.Markdown("", visible=False)##### 新增弹窗
            # =============================================================

        def initialize_system():
            """初始化系统状态"""
            initial_data = []
            if os.path.exists(output_path):
                try:
                    with open(output_path, 'r', encoding='utf-8') as f:
                        initial_data = json.load(f)
                        if not isinstance(initial_data, list):  # 强制类型校验
                            initial_data = []
                except Exception as e:
                    print(f"数据加载错误: {str(e)}")
                    initial_data = []
            
            cases = load_unprocessed_files(output_path, initial_data)
            if not cases:
                return [
                    {"id": "N/A", "transcript": "无待处理案件"}, 
                    [],
                    "当前进度：0/0",
                    gr.Column(visible=False),
                    *([False]*10),
                    *([""]*10),
                    "系统状态：所有案件已完成标注",
                    [],
                    initial_data
                ]
            
            first_case = cases[0]
            fact_count = first_case["fact_count"]  # 改动：获取事实数量
            facts = first_case["facts"] + [""]*(fact_count-len(first_case["facts"]))
            for i in range(fact_count):  # 新增
                with gr.Row():
                    fact_checkboxes.append(gr.Checkbox(label=f"事实 {i+1}", value=False))
                    fact_textboxes.append(gr.Textbox(interactive=True, lines=1, show_label=False))
            
            return [
                first_case,
                cases,
                f"当前进度：{len(cases)}",
                gr.Column(visible=True),
                *([False]*fact_count),
                *facts,
                "系统状态：请开始标注",
                [],
                initial_data
            ]

        def process_submission(current, cases, dataset, *components):
            """处理标注提交"""
            if not current or current["id"] == "N/A":
                return [current, cases, progress, fact_panel, *([False]*10), *([""]*10), status, [], dataset]
            
            # 类型安全处理
            if not isinstance(dataset, list):
                dataset = []
            # 改动：获取事实数量
            fact_count = current["fact_count"] if isinstance(current, dict) and "fact_count" in current else 10##新增
            # 解析组件状态#改变下面组件
            checkboxes = components[:fact_count]
            textboxes = components[fact_count:2*fact_count]
            custom_facts = components[2*fact_count] if len(components) > 2*fact_count else []

            # 收集选中事实
            selected = [tb.strip() for cb, tb in zip(checkboxes, textboxes) if cb and tb.strip()]
            
            # 处理自定义事实
            if isinstance(custom_facts, list):
                selected += [f.strip() for f in custom_facts if f.strip()]
            elif isinstance(custom_facts, str) and custom_facts.strip():
                selected.append(custom_facts.strip())

            # 创建新条目
            new_entry = {
                "卷宗编号": current["id"],
                "庭审笔录": current["transcript"],
                "事实标注": selected
            }
            
            # 安全更新数据
            updated_data = dataset.copy() if isinstance(dataset, list) else []
            updated_data.append(new_entry)
            
            # 保存到文件
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(updated_data, f, ensure_ascii=False, indent=4)
            
            # 获取下个案件
            remaining = cases[1:] if isinstance(cases, list) else []
            if not remaining:
                return [
                    {"id": "N/A", "transcript": "无待处理案件"},
                    [],
                    "当前进度：0/0",
                    gr.Column(visible=False),
                    *([False]*10),
                    *([""]*10),
                    "系统状态：标注完成",
                    [],
                    updated_data
                ]
            
            next_case = remaining[0]
            next_fact_count = next_case["fact_count"]  # 改动：获取下一个案件的事实数量
            next_facts = next_case["facts"] + [""]*(next_fact_count-len(next_case["facts"]))#改动：根据事实数量进行填充
            for i in range(fact_count):  # 新增
                with gr.Row():
                    fact_checkboxes.append(gr.Checkbox(label=f"事实 {i+1}", value=False))
                    fact_textboxes.append(gr.Textbox(interactive=True, lines=1, show_label=False))
            
            return [
                next_case,
                remaining,
                f"当前进度：{len(remaining)}",
                gr.Column(visible=True),
                *([False]*next_fact_count),
                *next_facts,
                "系统状态：请继续标注",
                [],
                updated_data
            ]

        def refresh_system(dataset):
            """刷新案件列表"""
            if not isinstance(dataset, list):
                dataset = []
            
            cases = load_unprocessed_files(output_path, dataset)
            if not cases:
                return [
                    {"id": "N/A", "transcript": "无待处理案件"}, 
                    [],
                    "当前进度：0/0",
                    gr.Column(visible=False),
                    *([False]*10),
                    *([""]*10),
                    "系统状态：所有案件已完成标注",
                    [],
                    dataset
                ]
            
            first = cases[0]
            fact_count = first["fact_count"]  # 改动：获取事实数量
            facts = first["facts"] + [""]*(fact_count-len(first["facts"]))# 改动，根据事实进行填充
            for i in range(fact_count):  # 新增
                with gr.Row():
                    fact_checkboxes.append(gr.Checkbox(label=f"事实 {i+1}", value=False))
                    fact_textboxes.append(gr.Textbox(interactive=True, lines=1, show_label=False))
            
            return [
                first,
                cases,
                f"当前进度：{len(cases)}",
                gr.Column(visible=True),
                *([False]*fact_count),
                *facts,
                "系统状态：请开始标注",
                [],
                dataset
            ]

        # 事件绑定
        interface.load(
            initialize_system,
            outputs=[current_case, case_queue, progress, fact_panel, 
                    *fact_checkboxes, *fact_textboxes, status, added_facts, dataset]
        )
        
        submit_btn.click(
            process_submission,
            inputs=[current_case, case_queue, dataset, *fact_checkboxes, *fact_textboxes, added_facts],
            outputs=[current_case, case_queue, progress, fact_panel, 
                    *fact_checkboxes, *fact_textboxes, status, added_facts, dataset]
        )
        
        refresh_btn.click(
            refresh_system,
            inputs=[dataset],
            outputs=[current_case, case_queue, progress, fact_panel,
                    *fact_checkboxes, *fact_textboxes, status, added_facts, dataset]
        )
        ################ 新增功能 ##########
        def show_then_hide_hint():
            # 显示提示框
            time.sleep(1)  # 等待1秒
            return gr.update(visible=False, value="")
        ##################### 新增功能 ##########
        add_btn.click(
            # lambda c, n, a: [c, (a if isinstance(a, list) else []) + [n.strip()]] if n.strip() else [c, a],
            # inputs=[current_case, new_fact_input, added_facts],
            # outputs=[current_case, added_facts]
            lambda c, n, a: [c, (a if isinstance(a, list) else []) + [n.strip()], gr.update(visible=True, value="事实新增成功！")] if n.strip() else [c, a, gr.update(visible=False, value="")],
            inputs=[current_case, new_fact_input, added_facts],
            outputs=[current_case, added_facts, hint]
        )
        hint.change(
            show_then_hide_hint,
            inputs=[],
            outputs=[hint],
            show_progress=False
        )
        ###############
        current_case.change(
            lambda x: [x["id"], x["transcript"]] if x else ["N/A", "无内容"],
            inputs=[current_case],
            outputs=[case_id, transcript]
        )

    return interface

if __name__ == "__main__":
    output_path = '/home/tmpuser/ltc/annotated_data.json'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    app = create_interface(output_path)
    app.launch(
        server_name="0.0.0.0",
        server_port=7860,
        show_error=True
    )